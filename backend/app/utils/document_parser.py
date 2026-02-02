"""
Document parsing utilities for PDF and DOCX files.
Extracts text, structure, and metadata from legal documents.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.docx import partition_docx
from unstructured.documents.elements import (
    Element,
    Title,
    NarrativeText,
    ListItem,
    Table,
)

logger = logging.getLogger(__name__)


class DocumentElement:
    """Represents a parsed document element with hierarchy."""

    def __init__(
        self,
        element_type: str,
        text: str,
        page_number: Optional[int] = None,
        hierarchy_level: int = 0,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        self.element_type = element_type
        self.text = text
        self.page_number = page_number
        self.hierarchy_level = hierarchy_level
        self.metadata = metadata or {}

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "element_type": self.element_type,
            "text": self.text,
            "page_number": self.page_number,
            "hierarchy_level": self.hierarchy_level,
            "metadata": self.metadata,
        }

    def __repr__(self) -> str:
        return f"DocumentElement(type={self.element_type}, page={self.page_number}, level={self.hierarchy_level})"


class DocumentParser:
    """Base document parser with common functionality."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_ext = Path(file_path).suffix.lower()
        self.elements: List[DocumentElement] = []
        self.metadata: Dict[str, Any] = {}

    def parse(self) -> List[DocumentElement]:
        """Parse document and return structured elements."""
        raise NotImplementedError("Subclasses must implement parse()")

    def extract_metadata(self) -> Dict[str, Any]:
        """Extract document metadata."""
        return {
            "filename": Path(self.file_path).name,
            "file_size": os.path.getsize(self.file_path),
            "file_type": self.file_ext,
        }


class PDFParser(DocumentParser):
    """PDF parser using PyMuPDF and Unstructured."""

    def __init__(self, file_path: str, use_ocr: bool = False):
        super().__init__(file_path)
        self.use_ocr = use_ocr

    def parse(self) -> List[DocumentElement]:
        """
        Parse PDF using Unstructured for structure detection.
        Falls back to PyMuPDF for simple text extraction if Unstructured fails.
        """
        logger.info(f"Parsing PDF: {self.file_path}")

        try:
            # Try Unstructured first for better structure detection
            elements = partition_pdf(
                filename=self.file_path,
                strategy="hi_res" if self.use_ocr else "fast",
                include_page_breaks=True,
                languages=["vie", "eng"],  # Vietnamese and English
            )

            self.elements = self._convert_unstructured_elements(elements)
            logger.info(f"Extracted {len(self.elements)} elements using Unstructured")

        except Exception as e:
            logger.warning(f"Unstructured parsing failed: {e}. Falling back to PyMuPDF")
            self.elements = self._parse_with_pymupdf()

        self.metadata = self._extract_pdf_metadata()
        return self.elements

    def _convert_unstructured_elements(self, elements: List[Element]) -> List[DocumentElement]:
        """Convert Unstructured elements to DocumentElement objects."""
        parsed_elements = []
        current_page = 1

        for elem in elements:
            # Determine hierarchy level based on element type
            if isinstance(elem, Title):
                hierarchy_level = self._infer_title_level(elem.text)
                element_type = f"title_level_{hierarchy_level}"
            elif isinstance(elem, ListItem):
                element_type = "list_item"
                hierarchy_level = 3
            elif isinstance(elem, Table):
                element_type = "table"
                hierarchy_level = 2
            else:
                element_type = "text"
                hierarchy_level = 4

            # Get page number from metadata
            page_number = getattr(elem.metadata, "page_number", current_page)
            if page_number:
                current_page = page_number

            parsed_elem = DocumentElement(
                element_type=element_type,
                text=elem.text,
                page_number=page_number,
                hierarchy_level=hierarchy_level,
                metadata={
                    "coordinates": getattr(elem.metadata, "coordinates", None),
                    "file_directory": getattr(elem.metadata, "file_directory", None),
                },
            )
            parsed_elements.append(parsed_elem)

        return parsed_elements

    def _infer_title_level(self, text: str) -> int:
        """
        Infer hierarchy level for Vietnamese legal document titles.
        Level 1: Document title
        Level 2: Chapters (Chương)
        Level 3: Sections (Mục)
        Level 4: Articles (Điều)
        Level 5: Clauses (Khoản)
        """
        text_lower = text.lower().strip()

        # Vietnamese legal structure patterns
        if any(keyword in text_lower for keyword in ["luật", "nghị định", "thông tư", "quyết định"]):
            return 1
        elif text_lower.startswith("chương"):
            return 2
        elif text_lower.startswith("mục"):
            return 3
        elif text_lower.startswith("điều"):
            return 4
        elif any(text_lower.startswith(str(i) + ".") for i in range(1, 100)):
            return 5
        else:
            return 3  # Default to section level

    def _parse_with_pymupdf(self) -> List[DocumentElement]:
        """Fallback parser using PyMuPDF for simple text extraction."""
        elements = []

        with fitz.open(self.file_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()

                # Split by paragraphs
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

                for para in paragraphs:
                    # Simple heuristic for structure
                    if self._is_title(para):
                        hierarchy_level = self._infer_title_level(para)
                        element_type = f"title_level_{hierarchy_level}"
                    else:
                        hierarchy_level = 4
                        element_type = "text"

                    elem = DocumentElement(
                        element_type=element_type,
                        text=para,
                        page_number=page_num,
                        hierarchy_level=hierarchy_level,
                    )
                    elements.append(elem)

        return elements

    def _is_title(self, text: str) -> bool:
        """Heuristic to detect if text is a title/heading."""
        text_lower = text.lower().strip()
        return (
            len(text) < 200
            and (
                text_lower.startswith(("chương", "mục", "điều", "luật", "nghị định"))
                or text.isupper()
            )
        )

    def _extract_pdf_metadata(self) -> Dict[str, Any]:
        """Extract metadata from PDF."""
        metadata = self.extract_metadata()

        try:
            with fitz.open(self.file_path) as doc:
                metadata.update({
                    "page_count": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creation_date": doc.metadata.get("creationDate", ""),
                    "modification_date": doc.metadata.get("modDate", ""),
                })
        except Exception as e:
            logger.error(f"Failed to extract PDF metadata: {e}")

        return metadata


class DOCXParser(DocumentParser):
    """DOCX parser using python-docx and Unstructured."""

    def parse(self) -> List[DocumentElement]:
        """Parse DOCX file and extract structured elements."""
        logger.info(f"Parsing DOCX: {self.file_path}")

        try:
            # Try Unstructured first
            elements = partition_docx(filename=self.file_path)
            self.elements = self._convert_unstructured_elements(elements)
            logger.info(f"Extracted {len(self.elements)} elements using Unstructured")

        except Exception as e:
            logger.warning(f"Unstructured parsing failed: {e}. Falling back to python-docx")
            self.elements = self._parse_with_docx()

        self.metadata = self._extract_docx_metadata()
        return self.elements

    def _convert_unstructured_elements(self, elements: List[Element]) -> List[DocumentElement]:
        """Convert Unstructured elements to DocumentElement objects."""
        parsed_elements = []

        for elem in elements:
            if isinstance(elem, Title):
                hierarchy_level = self._infer_title_level(elem.text)
                element_type = f"title_level_{hierarchy_level}"
            elif isinstance(elem, ListItem):
                element_type = "list_item"
                hierarchy_level = 3
            elif isinstance(elem, Table):
                element_type = "table"
                hierarchy_level = 2
            else:
                element_type = "text"
                hierarchy_level = 4

            parsed_elem = DocumentElement(
                element_type=element_type,
                text=elem.text,
                page_number=None,  # DOCX doesn't have page numbers reliably
                hierarchy_level=hierarchy_level,
            )
            parsed_elements.append(parsed_elem)

        return parsed_elements

    def _infer_title_level(self, text: str) -> int:
        """Infer hierarchy level for Vietnamese legal document titles."""
        text_lower = text.lower().strip()

        if any(keyword in text_lower for keyword in ["luật", "nghị định", "thông tư"]):
            return 1
        elif text_lower.startswith("chương"):
            return 2
        elif text_lower.startswith("mục"):
            return 3
        elif text_lower.startswith("điều"):
            return 4
        else:
            return 3

    def _parse_with_docx(self) -> List[DocumentElement]:
        """Fallback parser using python-docx."""
        elements = []
        doc = DocxDocument(self.file_path)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect hierarchy from style or content
            style_name = para.style.name.lower()

            if "heading" in style_name or "title" in style_name:
                hierarchy_level = self._infer_title_level(text)
                element_type = f"title_level_{hierarchy_level}"
            else:
                hierarchy_level = 4
                element_type = "text"

            elem = DocumentElement(
                element_type=element_type,
                text=text,
                page_number=None,
                hierarchy_level=hierarchy_level,
            )
            elements.append(elem)

        # Parse tables
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            elem = DocumentElement(
                element_type="table",
                text=table_text,
                page_number=None,
                hierarchy_level=2,
            )
            elements.append(elem)

        return elements

    def _extract_table_text(self, table) -> str:
        """Extract text from DOCX table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _extract_docx_metadata(self) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = self.extract_metadata()

        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties

            metadata.update({
                "page_count": len(doc.paragraphs),  # Approximation
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "creation_date": core_props.created.isoformat() if core_props.created else "",
                "modification_date": core_props.modified.isoformat() if core_props.modified else "",
            })
        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {e}")

        return metadata


def parse_document(file_path: str, use_ocr: bool = False) -> tuple[List[DocumentElement], Dict[str, Any]]:
    """
    Factory function to parse a document based on file extension.

    Args:
        file_path: Path to the document file
        use_ocr: Whether to use OCR for PDFs (slower but better for scanned docs)

    Returns:
        Tuple of (elements, metadata)
    """
    file_ext = Path(file_path).suffix.lower()

    if file_ext == ".pdf":
        parser = PDFParser(file_path, use_ocr=use_ocr)
    elif file_ext in [".docx", ".doc"]:
        parser = DOCXParser(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

    elements = parser.parse()
    metadata = parser.metadata

    return elements, metadata
